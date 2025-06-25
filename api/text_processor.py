
import logging
from pathlib import Path
import pandas as pd
import io
from typing import Dict, Any, Optional, Tuple, List
import re

# Import PDF processing libraries
try:
    import pdfplumber
    PDF_PROCESSING_AVAILABLE = True
except ImportError:
    PDF_PROCESSING_AVAILABLE = False
    logging.warning("pdfplumber not available. PDF processing will be limited.")

try:
    import docx
    DOCX_PROCESSING_AVAILABLE = True
except ImportError:
    DOCX_PROCESSING_AVAILABLE = False
    logging.warning("python-docx not available. DOCX processing will be limited.")

try:
    import openpyxl
    EXCEL_PROCESSING_AVAILABLE = True
except ImportError:
    EXCEL_PROCESSING_AVAILABLE = False
    logging.warning("openpyxl not available. Excel processing will be limited.")

logger = logging.getLogger(__name__)

def count_words(text: str) -> int:
    """Counts words in a given string."""
    if not text:
        return 0
    return len(text.split())

def extract_text_from_pdf(file_path: str) -> Tuple[Optional[str], int]:
    """Extracts text from a PDF file using pdfplumber."""
    if not PDF_PROCESSING_AVAILABLE:
        logger.warning("PDF processing library not available")
        return None, 0
    
    text_content = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                    continue
        
        full_text = "\n\n".join(text_content)
        word_count = count_words(full_text)
        
        if not full_text.strip():
            # Try to extract basic info if no text found
            full_text = f"PDF document with {len(pdf.pages)} pages. Text extraction may require OCR for scanned documents."
            word_count = count_words(full_text)
            
        return full_text, word_count
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {e}")
        return None, 0

def extract_text_from_docx(file_path: str) -> Tuple[Optional[str], int]:
    """Extracts text from a DOCX file."""
    if not DOCX_PROCESSING_AVAILABLE:
        logger.warning("DOCX processing library not available")
        return None, 0
    
    try:
        doc = docx.Document(file_path)
        text_content = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_text.append(" | ".join(row_text))
            if table_text:
                text_content.append("--- Table ---\n" + "\n".join(table_text))
        
        full_text = "\n\n".join(text_content)
        return full_text, count_words(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {e}")
        return None, 0

def extract_text_from_txt(file_path: str) -> Tuple[Optional[str], int]:
    """Extracts text from a TXT file."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            full_text = f.read()
        return full_text, count_words(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from TXT {file_path}: {e}")
        return None, 0

def extract_text_from_csv(file_path: str) -> Tuple[Optional[str], int]:
    """Extracts text from a CSV file by concatenating all string columns."""
    try:
        df = pd.read_csv(file_path, dtype=str, keep_default_na=False)
        text_content = []
        
        # Add headers
        text_content.append("--- CSV Headers ---")
        text_content.append(" | ".join(df.columns.tolist()))
        text_content.append("")
        
        # Add data summary
        text_content.append("--- Data Summary ---")
        text_content.append(f"Rows: {len(df)}, Columns: {len(df.columns)}")
        text_content.append("")
        
        # Add sample data
        text_content.append("--- Sample Data ---")
        for idx, row in df.head(10).iterrows():
            row_text = []
            for col in df.columns:
                cell_value = str(row[col]).strip()
                if cell_value and cell_value != 'nan':
                    row_text.append(f"{col}: {cell_value}")
            if row_text:
                text_content.append(" | ".join(row_text))
        
        # Add all text data for analysis
        text_content.append("\n--- All Text Data ---")
        for col in df.columns:
            col_text = df[col].astype(str).str.strip()
            col_text = col_text[col_text != ''].dropna()
            if not col_text.empty:
                text_content.append(f"{col}: {' '.join(col_text.tolist())}")
        
        full_text = "\n".join(text_content)
        return full_text, count_words(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from CSV {file_path}: {e}")
        return None, 0

def extract_text_from_excel(file_path: str) -> Tuple[Optional[str], int]:
    """Extracts text from an XLSX file by concatenating all string cells from all sheets."""
    if not EXCEL_PROCESSING_AVAILABLE:
        logger.warning("Excel processing library not available")
        return None, 0
    
    try:
        xls = pd.ExcelFile(file_path)
        all_sheets_text = []
        
        for sheet_name in xls.sheet_names:
            try:
                df = pd.read_excel(xls, sheet_name=sheet_name, dtype=str, keep_default_na=False)
                sheet_text_content = []
                
                sheet_text_content.append(f"=== Sheet: {sheet_name} ===")
                sheet_text_content.append(f"Dimensions: {len(df)} rows x {len(df.columns)} columns")
                
                # Add headers
                sheet_text_content.append("Headers: " + " | ".join(df.columns.tolist()))
                
                # Add sample data
                for idx, row in df.head(5).iterrows():
                    row_text = []
                    for col in df.columns:
                        cell_value = str(row[col]).strip()
                        if cell_value and cell_value != 'nan':
                            row_text.append(f"{col}: {cell_value}")
                    if row_text:
                        sheet_text_content.append(" | ".join(row_text))
                
                # Add all text data
                for col in df.columns:
                    col_text = df[col].astype(str).str.strip()
                    col_text = col_text[col_text != ''].dropna()
                    if not col_text.empty:
                        sheet_text_content.append(f"{col} data: {' '.join(col_text.tolist()[:100])}")  # Limit to 100 items
                
                all_sheets_text.append("\n".join(sheet_text_content))
            except Exception as e:
                logger.warning(f"Error processing sheet {sheet_name}: {e}")
                continue

        full_text = "\n\n".join(all_sheets_text)
        return full_text, count_words(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from Excel {file_path}: {e}")
        return None, 0

def extract_text_from_file(file_path_str: str, file_extension: str, mime_type: Optional[str] = None) -> Optional[Dict[str, Any]]:
    """
    Main function to extract text from a file based on its extension or mime type.
    Returns a dictionary with 'text' and 'word_count' or None if extraction fails.
    """
    file_path = Path(file_path_str)
    if not file_path.exists():
        logger.error(f"File not found for extraction: {file_path_str}")
        return None

    extracted_text = None
    word_count = 0
    metadata = {}

    logger.info(f"Attempting to extract text from {file_path_str} with extension {file_extension}")

    try:
        if file_extension.lower() == ".pdf":
            extracted_text, word_count = extract_text_from_pdf(file_path_str)
            metadata['processing_method'] = 'pdfplumber' if PDF_PROCESSING_AVAILABLE else 'basic'
        elif file_extension.lower() == ".docx":
            extracted_text, word_count = extract_text_from_docx(file_path_str)
            metadata['processing_method'] = 'python-docx' if DOCX_PROCESSING_AVAILABLE else 'basic'
        elif file_extension.lower() in [".txt", ".md"]:
            extracted_text, word_count = extract_text_from_txt(file_path_str)
            metadata['processing_method'] = 'text_reader'
        elif file_extension.lower() == ".csv":
            extracted_text, word_count = extract_text_from_csv(file_path_str)
            metadata['processing_method'] = 'pandas_csv'
        elif file_extension.lower() in [".xlsx", ".xls"]:
            extracted_text, word_count = extract_text_from_excel(file_path_str)
            metadata['processing_method'] = 'pandas_excel' if EXCEL_PROCESSING_AVAILABLE else 'basic'
        else:
            logger.warning(f"Unsupported file extension for text extraction: {file_extension} for file {file_path_str}")
            return None

        if extracted_text is not None:
            # Additional metadata
            metadata.update({
                'file_size_bytes': file_path.stat().st_size,
                'character_count': len(extracted_text),
                'line_count': len(extracted_text.split('\n')),
                'file_extension': file_extension,
                'processing_status': 'success'
            })
            
            logger.info(f"Successfully extracted {word_count} words from {file_path_str}.")
            return {
                "text": extracted_text, 
                "word_count": word_count,
                "metadata": metadata
            }
        else:
            logger.error(f"Failed to extract text from {file_path_str}.")
            return None
    except Exception as e:
        logger.error(f"Unexpected error during text extraction from {file_path_str}: {e}")
        return None

# Enhanced text analysis
BUSINESS_KEYWORDS = {
    "market_trends": ["market trends", "market trend", "trending", "emerging market", "market growth", "market expansion"],
    "competitive_landscape": ["competitive landscape", "competition", "competitors", "market share", "competitive advantage", "market position"],
    "customer_behavior": ["customer behavior", "customer behaviour", "user behavior", "consumer behavior", "customer journey", "customer experience"],
    "growth_strategy": ["growth strategy", "strategic growth", "expansion strategy", "business growth", "scaling", "market penetration"],
    "financial_performance": ["revenue", "profit", "financial performance", "ROI", "return on investment", "financial metrics"],
    "technology_trends": ["AI", "artificial intelligence", "machine learning", "blockchain", "digital transformation", "automation"],
    "market_analysis": ["market analysis", "market research", "market intelligence", "market insights", "market data"],
    "business_strategy": ["business strategy", "strategic planning", "strategic initiatives", "business model", "value proposition"]
}

def analyze_text_keywords(text_content: str) -> Dict[str, Any]:
    """
    Analyzes the text content for the presence of predefined business keywords.
    Performs a case-insensitive search and provides detailed analysis.
    """
    if not text_content:
        return {
            "keyword_categories": {category: False for category in BUSINESS_KEYWORDS.keys()},
            "keyword_matches": {},
            "analysis_summary": "No text content provided for analysis",
            "text_quality": "poor"
        }

    lower_text_content = text_content.lower()
    keyword_matches = {}
    category_results = {}
    
    for category, keywords in BUSINESS_KEYWORDS.items():
        matches = []
        for keyword in keywords:
            if keyword.lower() in lower_text_content:
                # Count occurrences
                count = lower_text_content.count(keyword.lower())
                matches.append({"keyword": keyword, "count": count})
        
        category_results[category] = len(matches) > 0
        if matches:
            keyword_matches[category] = matches

    # Calculate text quality score
    total_categories = len(BUSINESS_KEYWORDS)
    matched_categories = sum(1 for result in category_results.values() if result)
    quality_score = matched_categories / total_categories
    
    if quality_score >= 0.7:
        text_quality = "excellent"
    elif quality_score >= 0.5:
        text_quality = "good"
    elif quality_score >= 0.3:
        text_quality = "fair"
    else:
        text_quality = "poor"

    # Generate summary
    word_count = count_words(text_content)
    char_count = len(text_content)
    
    analysis_summary = f"Analyzed {word_count} words, {char_count} characters. "
    analysis_summary += f"Found business-relevant content in {matched_categories}/{total_categories} categories. "
    analysis_summary += f"Text quality: {text_quality}."

    logger.info(f"Keyword analysis complete. Quality: {text_quality}, Categories matched: {matched_categories}/{total_categories}")
    
    return {
        "keyword_categories": category_results,
        "keyword_matches": keyword_matches,
        "analysis_summary": analysis_summary,
        "text_quality": text_quality,
        "quality_score": quality_score,
        "matched_categories_count": matched_categories,
        "total_categories": total_categories,
        "word_count": word_count,
        "character_count": char_count
    }

def extract_entities(text_content: str) -> Dict[str, List[str]]:
    """
    Extract basic entities from text using regex patterns.
    In production, this could be enhanced with NLP libraries like spaCy.
    """
    entities = {
        "emails": [],
        "urls": [],
        "phone_numbers": [],
        "dates": [],
        "companies": [],
        "currencies": []
    }
    
    if not text_content:
        return entities
    
    # Email pattern
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    entities["emails"] = list(set(re.findall(email_pattern, text_content)))
    
    # URL pattern
    url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
    entities["urls"] = list(set(re.findall(url_pattern, text_content)))
    
    # Phone number pattern (basic)
    phone_pattern = r'\b(?:\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}\b'
    entities["phone_numbers"] = list(set(re.findall(phone_pattern, text_content)))
    
    # Date pattern (basic)
    date_pattern = r'\b(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}[/-]\d{1,2}[/-]\d{1,2})\b'
    entities["dates"] = list(set(re.findall(date_pattern, text_content)))
    
    # Currency pattern
    currency_pattern = r'\$\d+(?:,\d{3})*(?:\.\d{2})?|\b\d+(?:,\d{3})*(?:\.\d{2})?\s*(?:USD|EUR|GBP|JPY)\b'
    entities["currencies"] = list(set(re.findall(currency_pattern, text_content)))
    
    return entities

# Test functionality
if __name__ == '__main__':
    # Test with sample text
    sample_text = """
    This document discusses market trends and the competitive landscape.
    Understanding customer behavior is key for our new growth strategy.
    We need to analyze financial performance and ROI metrics.
    AI and machine learning are driving digital transformation.
    Contact us at info@company.com or visit https://company.com
    Revenue increased to $1,250,000 in Q1 2024.
    """
    
    print("Testing Keyword Analysis:")
    analysis_result = analyze_text_keywords(sample_text)
    print(f"Analysis: {analysis_result}")
    print()
    
    print("Testing Entity Extraction:")
    entities = extract_entities(sample_text)
    print(f"Entities: {entities}")
